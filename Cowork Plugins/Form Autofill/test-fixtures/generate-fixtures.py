"""Generates test fixtures for the Form Autofill Cowork plugin.

Run:  python generate-fixtures.py

Produces three attachments that together exercise every branch of the skill:
  new-starter-form.docx   fillable, all three tiers, deliberate synonym variants
  travel-profile.xlsx     fillable, overlapping fields to test ask-once dedupe
  benefits-guide.docx     reference material that must NOT be filled
"""

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.worksheet.datavalidation import DataValidation

HERE = os.path.dirname(os.path.abspath(__file__))
BLANK = "_" * 42


def _field_table(doc, rows):
    """Label/value table. Value cells are empty on purpose."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, label in enumerate(rows):
        cell = table.cell(i, 0)
        cell.text = label
        cell.paragraphs[0].runs[0].font.bold = True
        cell.width = Pt(200)
        table.cell(i, 1).text = ""
    doc.add_paragraph()
    return table


def _blank_lines(doc, labels):
    for label in labels:
        doc.add_paragraph(f"{label}: {BLANK}")
    doc.add_paragraph()


def build_word_form(path):
    doc = Document()

    doc.add_heading("Contoso Group - New Starter Onboarding Form", level=0)
    p = doc.add_paragraph(
        "Please complete all sections and return this form to HR Shared Services "
        "within five working days of your start date. Sections marked * are required."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Tier 1 - resolvable from the directory profile.
    # Labels deliberately use synonyms, not the canonical field names.
    doc.add_heading("Section A - Employee Details *", level=1)
    _field_table(doc, [
        "Surname",
        "Given Name",
        "Position",
        "Division",
        "Work Site",
        "Business Email",
        "Staff Number",
        "Reports To",
    ])

    # Tier 2 - interview, vault-storable. Rendered as underscore lines
    # rather than a table so both fill modes get exercised.
    doc.add_heading("Section B - Personal Details *", level=1)
    _blank_lines(doc, [
        "Name as it appears on your ID",
        "Birthdate",
        "Residential Address",
        "Cell",
        "Alternate Email",
    ])

    doc.add_heading("Section C - In Case of Emergency *", level=1)
    _field_table(doc, [
        "ICE Contact",
        "Relation to Employee",
        "Contact Number",
    ])

    # Tier 3 - must never be auto-filled or stored.
    doc.add_heading("Section D - Payroll and Compliance *", level=1)
    doc.add_paragraph(
        "This information is required for payroll processing and right-to-work checks."
    )
    _field_table(doc, [
        "National Insurance Number",
        "Bank Account Number",
        "Sort Code",
        "Passport No",
    ])

    doc.add_heading("Section E - Preferences (optional)", level=1)
    _blank_lines(doc, [
        "Meal Preference",
        "Apparel Size",
        "Accessibility Requirements",
    ])

    # Signature is Tier 3: an act of authorization, never reproduced.
    doc.add_heading("Section F - Declaration *", level=1)
    doc.add_paragraph(
        "I confirm the information given above is accurate and complete."
    )
    _blank_lines(doc, ["Authorised Signature", "Date"])

    doc.save(path)


def build_excel_form(path):
    wb = Workbook()

    ws = wb.active
    ws.title = "Traveller Profile"

    header = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="0078D4")
    label_font = Font(bold=True)
    locked_fill = PatternFill("solid", fgColor="F3F2F1")

    ws["A1"] = "Contoso Travel - Traveller Profile"
    ws["A1"].font = Font(bold=True, size=14)
    ws.merge_cells("A1:B1")

    ws["A2"] = "Field"
    ws["B2"] = "Your Answer"
    for c in ("A2", "B2"):
        ws[c].font = header
        ws[c].fill = header_fill

    # Several of these intentionally duplicate the Word form under different
    # labels - the skill should ask for each concept once, not twice.
    fields = [
        "Full Legal Name",
        "Date of Birth",
        "Home Address",
        "Mobile Number",
        "Passport Number",
        "Passport Expiry",
        "Emergency Contact Name",
        "Emergency Contact Number",
        "Dietary Requirements",
        "Frequent Flyer Number",
    ]
    for i, label in enumerate(fields, start=3):
        ws[f"A{i}"] = label
        ws[f"A{i}"].font = label_font
        ws[f"A{i}"].fill = locked_fill
        ws[f"B{i}"] = None

    last = 2 + len(fields)

    # A calculated cell. Must not be written into.
    ws[f"A{last + 2}"] = "Fields completed (auto)"
    ws[f"A{last + 2}"].font = label_font
    ws[f"A{last + 2}"].fill = locked_fill
    ws[f"B{last + 2}"] = f"=COUNTA(B3:B{last})"

    dv = DataValidation(
        type="list",
        formula1='"None,Vegetarian,Vegan,Halal,Kosher,Gluten-free"',
        allow_blank=True,
    )
    ws.add_data_validation(dv)
    dietary_row = 3 + fields.index("Dietary Requirements")
    dv.add(ws[f"B{dietary_row}"])

    ws.column_dimensions["A"].width = 30
    ws.column_dimensions["B"].width = 42
    for row in ws.iter_rows(min_row=3, max_row=last, min_col=2, max_col=2):
        for cell in row:
            cell.alignment = Alignment(vertical="center")

    # Second sheet is reference material, not input.
    info = wb.create_sheet("Instructions")
    info["A1"] = "How to complete this form"
    info["A1"].font = Font(bold=True, size=13)
    lines = [
        "",
        "1. Complete every field on the 'Traveller Profile' sheet.",
        "2. Passport details must match your travel document exactly.",
        "3. The 'Fields completed' cell calculates automatically - do not edit it.",
        "4. Return the workbook to travel@contoso.example.",
        "",
        "Questions? Contact the travel desk on extension 4400.",
    ]
    for i, line in enumerate(lines, start=2):
        info[f"A{i}"] = line
    info.column_dimensions["A"].width = 70

    wb.save(path)


def build_reference_doc(path):
    """Prose only, no blanks. Must be reported and left untouched."""
    doc = Document()
    doc.add_heading("Contoso Group - Employee Benefits Guide", level=0)
    doc.add_paragraph(
        "This guide summarises the benefits available to Contoso Group employees. "
        "It is provided for information only. You do not need to return this document."
    )

    doc.add_heading("Pension", level=1)
    doc.add_paragraph(
        "Contoso operates a defined contribution pension scheme. Employees are "
        "enrolled automatically after three months of continuous service. The "
        "employer contribution is 6% of pensionable pay, matched up to a further 3%."
    )

    doc.add_heading("Private Medical Cover", level=1)
    doc.add_paragraph(
        "Cover is provided through Contoso's group scheme and includes outpatient "
        "treatment, diagnostics, and mental health support. Partners and dependants "
        "may be added during the annual enrolment window each October."
    )

    doc.add_heading("Annual Leave", level=1)
    doc.add_paragraph(
        "Full-time employees receive 25 days of annual leave plus public holidays, "
        "rising to 28 days after five years of service. Up to five days may be "
        "carried into the following leave year."
    )

    doc.add_heading("Learning and Development", level=1)
    doc.add_paragraph(
        "Each employee has an annual development allowance of 1,500 GBP for "
        "training, conferences, and professional memberships. Requests are approved "
        "by your line manager."
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "For questions about any benefit, contact HR Shared Services at "
        "hr@contoso.example."
    )

    doc.save(path)


def main():
    targets = [
        ("new-starter-form.docx", build_word_form),
        ("travel-profile.xlsx", build_excel_form),
        ("benefits-guide.docx", build_reference_doc),
    ]
    for name, builder in targets:
        path = os.path.join(HERE, name)
        builder(path)
        print(f"wrote {name} ({os.path.getsize(path):,} bytes)")


if __name__ == "__main__":
    main()
